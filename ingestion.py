from dotenv import load_dotenv
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
import pandas as pd
import json
from langchain_openai import OpenAIEmbeddings
from langchain_pinecone import PineconeVectorStore
from docx import Document as DocxDocument
import pandas as pd
import os

embeddings = OpenAIEmbeddings()


def load_word_documents_with_docx(folder_path):
    all_documents = []
    
    for file_name in os.listdir(folder_path):
        if file_name.endswith(".docx"):
            file_path = os.path.join(folder_path, file_name)
            
            try:
                # Load Word document content
                docx = DocxDocument(file_path)
                full_text = []
                for para in docx.paragraphs:
                    full_text.append(para.text)
                page_content = "\n".join(full_text)
                
                # Create LangChain Document
                all_documents.append(
                    Document(
                        page_content=page_content,
                        metadata={"file_type": "docx", "source": file_name}
                    )
                )
            except Exception as e:
                print(f"Error loading {file_name}: {e}")
    
    print(f"Loaded {len(all_documents)} documents from Word files in {folder_path}.")
    return all_documents

def load_json_documents(file_paths):
    """
    Load and process JSON files into a list of Document objects.

    Args:
        file_paths (list of str): List of JSON file paths to process.

    Returns:
        list of Document: A list of Document objects with content and metadata.
    """
    combined_json_documents = []

    for file_path in file_paths:
        try:
            # Read JSON file
            data = pd.read_json(file_path)
            
            # Convert each row in the JSON to a Document
            for _, row in data.iterrows():
                combined_json_documents.append(
                    Document(
                        page_content=str(row.to_dict()),
                        metadata={"file_type": "json", "source": file_path.split('/')[-1]}  # Add source dynamically
                    )
                )
        except ValueError as e:
            print(f"Error processing JSON file {file_path}: {e}")

    print(f"Loaded {len(combined_json_documents)} documents from JSON files.")
    return combined_json_documents



def ingest_docs():
    # Load PDF file
    pdf_loader = PyPDFLoader("./Onasi_RCM.pdf")
    pdf_documents = pdf_loader.load()
    for doc in pdf_documents:
        doc.metadata.update({"file_type": "pdf", "source": "Onasi_RCM.pdf"})
    print(f"Loaded {len(pdf_documents)} documents from PDF.")
    
    # Example usage
    folder_path = "./DHIS_documents"  # Path to the folder containing Word docs
    word_documents = load_word_documents_with_docx(folder_path)
    # Output metadata and content for verification
    for doc in word_documents[:5]:
        print(f"Metadata: {doc.metadata}, Content: {doc.page_content[:100]}")
    
    # Load JSON Files
    file_paths = ["./Medical_coding.json", "./Nphies_validation.json"]
    json_documents = load_json_documents(file_paths)

    # Output metadata and content for verification
    for doc in json_documents[:5]:  # Display first 5 documents
        print(f"Metadata: {doc.metadata}, Content: {doc.page_content[:100]}")
    
    # Combine all documents
    all_documents = pdf_documents + json_documents

    # Split documents into chunks for embedding, using specified chunk size and overlap
    text_splitter  = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=400)

    split_documents = []
    for doc in all_documents:
        if len(doc.page_content) > 500:
            split_documents.extend(text_splitter.split_documents([doc]))
        else:
            split_documents.append(doc)

    print(f"Going to add {len(split_documents)} to Pinecone")
    for doc in split_documents:
        print(f"Document Metadata: {doc.metadata}, Content Length: {len(doc.page_content)}")

    # Index documents in Pinecone
    PineconeVectorStore.from_documents(
        split_documents, embeddings, index_name="rcm-final-app"
    )
    print("****Loading to vectorstore done ***")

if __name__ == "__main__":
    ingest_docs()
